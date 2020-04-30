import re
from docx import Document
from docx.shared import Inches,RGBColor
document = Document()
document.add_heading('Python 运行日志', 0)
p = document.add_paragraph('运行正常')
a = '\033[37;44m执行时间--------------------------------------------\033[0m'
b = '11111'
c = '22222'
style = document.styles.add_style('aaaa',1)
style.font.color.rgb = RGBColor(255,0,0)
print(type(a))
for i in (b,c):
    print(i)
    if re.search('^22', i.strip().lower()):
        p = document.add_paragraph(i, style=style)
    else:
        p = document.add_paragraph(i)

document.save('Python 运行日志.docx')